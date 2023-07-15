#This code is used to extract our own requirements from HLR RAR

import docx
import os
def extract_proposed_statement(file_path):
    # Load the word document
    doc = docx.Document(file_path)
    
    # Iterate over the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph text contains the word "PROPOSED STATEMENT"
        if "PROPOSED STATEMENT" in paragraph.text:
            # Return the next paragraph as the proposed statement
            return paragraph.next_sibling.text

# Example usage
#i want to access every table inside this word file and extract the data from it
important_text=[]
for filename in os.listdir():
    if filename.endswith(".docx"):
        doc = docx.Document(filename)
        elements_array=[]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        for inner_table in cell.tables:
                            for inner_row in inner_table.rows:
                                for inner_cell in inner_row.cells:
                                    elements_array.append(inner_cell.text)
        print(elements_array[11],elements_array[-1])

        important_text.append(elements_array[11])
        important_text.append(elements_array[-1])
important_text=[x for x in important_text if x != "N/A"]
print(important_text)
with open("statementlar.txt", "w",encoding="utf-8") as f:
    for text in important_text:
        f.write(text + "\n")
#proposed_statement = extract_proposed_statement(file_path)
#print("Proposed Statement:", proposed_statement)